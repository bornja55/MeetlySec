"""
build_minutes_template.py — สร้างไฟล์ template ทั้งหมดใน `templates/` (Module 4, 2026-08-03 + แก้
เพิ่ม multi-template รองรับหลาย form/template ในวันเดียวกัน หลังผู้ใช้ถามว่าจะรองรับการประชุมที่ต้อง
ใช้ form อื่นได้อย่างไร)

**เหตุผลที่ต้องมีสคริปต์นี้แยกต่างหาก (ไม่ commit แค่ไฟล์ .docx เฉยๆ)**: เปิดไฟล์จริง
`260628 Draft_EMPIRE - BOD Minutes 15-2569 v.5.docx` ด้วย python-docx แล้วยืนยันว่าเป็นรายงานที่
เขียนเสร็จสมบูรณ์แล้ว **ไม่มี placeholder เลยสักจุด** — ถาม `AskUserQuestion` ก่อนตัดสินใจว่าจะเอาไฟล์
นี้มาใช้ตรงๆ ไม่ได้ (ไม่ generic พอ) เลยเลือก **สร้าง template ใหม่ทั้งหมดด้วย `docxtpl`** (Jinja tag
ฝังใน .docx) เลียนแบบ layout/ฟอนต์/การจัดหน้าโดยรวมจากไฟล์จริง (ฟอนต์ TH SarabunPSK 15pt, A4, margin
1 นิ้ว, จัดกึ่งกลางหัวเอกสาร, จัดพิสูจน์อักษรแบบ Thai-Justify) แต่ **เว้นตารางธุรกรรมไว้ให้ Maker เพิ่ม
เองใน Word หลัง generate** (ตัดสินใจคู่กับ Module 3's schema แบบยืดหยุ่น — ดู minutes_schema.py หัวไฟล์)

รันสคริปต์นี้ครั้งเดียว (`python build_minutes_template.py`) เพื่อ**สร้างไฟล์ template ทั้งหมดที่ commit
เก็บไว้ใน repo** — โค้ด runtime (`docx_generation.py`) แค่โหลด template ตาม
`docx_generation.TEMPLATE_REGISTRY` มา render ด้วย `docxtpl` ไม่ได้เรียกสคริปต์นี้ทุกครั้งที่สร้าง
เอกสาร เก็บสคริปต์นี้ไว้เผื่อต้องแก้ layout/เพิ่ม field ใหม่ในอนาคต (แก้ script นี้ → รันใหม่ → commit
ไฟล์ .docx ที่ได้ทับของเดิม) แทนที่จะแก้ XML ของ .docx ตรงๆด้วยมือ

**Multi-template**: `_build_document()` คือแกนกลางที่ใช้ร่วมกันทุก template (layout/ฟอนต์/Jinja tag
เหมือนกันหมด) รับพารามิเตอร์แค่ข้อความหัวเรื่อง (`doc_title_prefix`) ให้ต่างกันได้ — **template ใหม่
ทุกอันต้องใช้ Jinja tag ชุดเดียวกันเป๊ะ** (ดูรายการทั้งหมดด้านล่าง) เพราะข้อมูลมาจาก `minutes_json`
(Module 3's schema) ที่ไม่เปลี่ยนตาม template — ถ้าต้องการ field ใหม่ที่ไม่มีในนี้ ต้องแก้
`docx_generation.py`/`minutes_schema.py` เพิ่มด้วย ไม่ใช่แค่เพิ่มไฟล์ template

**วิธีเพิ่ม template ใหม่ในอนาคต (ไม่ต้องแก้ script นี้เลยก็ได้)**: วิธีที่ง่ายที่สุดคือก็อปปี้ไฟล์
`.docx` ที่มีอยู่แล้วไปเป็นชื่อใหม่ในโฟลเดอร์เดียวกัน เปิดแก้ layout/ถ้อยคำด้วย Microsoft Word ตรงๆ
(ห้ามลบ/พิมพ์ผิด Jinja tag เดิม) แล้วเพิ่ม 1 entry ใน `docx_generation.TEMPLATE_REGISTRY` ชี้ไปที่ไฟล์
ใหม่ — ไม่จำเป็นต้องแก้/รัน script นี้เลยถ้าไม่ได้เปลี่ยนโครงสร้างใหญ่

**Jinja tag ที่ต้องตรงกับ context ใน `docx_generation.py::render_minutes_docx()`**: company_name,
company_address, meeting_number, meeting_date_thai, attendees (list ของ {name, position}),
chairperson_name, agenda_items (list ของ {agenda_order, description, discussion_summary,
resolution_status_label, resolution_text}), other_business_notes, generated_by_model,
generated_at_thai

**หมายเหตุเรื่อง docxtpl's `{%p ... %}` tag**: ใช้แทน `{% ... %}` ธรรมดาตรงจุดที่เป็น for-loop ที่อยู่
ในพารากราฟของตัวเอง — docxtpl จะลบพารากราฟของ tag นั้นทิ้งทั้งพารากราฟหลัง render (ไม่ทิ้งบรรทัดว่าง
เปล่าค้างไว้เหมือน `{% %}` ธรรมดา) ต้องเป็นพารากราฟที่มีแค่ tag นี้ล้วนๆเท่านั้น (สร้างเป็น run เดียว
ไม่ปนกับข้อความอื่นในพารากราฟเดียวกัน — ไม่งั้น docxtpl parse ผิด)

**ยังไม่ได้ทำ/ข้อจำกัดที่รู้อยู่แล้ว**:
- ไม่ได้ใส่สี EMPIRE CI (Gold/Teal) ลงในเอกสารพวกนี้โดยตั้งใจ — เอกสารรายงานการประชุมเป็นเอกสารทาง
  กฎหมาย/compliance ต้นฉบับจริงที่เปิดดูเป็นขาว-ดำล้วน (ต่างจากดีไซน์ dashboard ที่บังคับ CI เต็มรูปแบบ)
- ไม่มีเลของค์ประชุม/ร้อยละองค์ประชุม/เวลาปิดประชุม/ชื่อผู้บันทึกรายงาน (Module 3's schema ไม่สร้าง
  ตัวเลข/ชื่อพวกนี้ เพื่อกันหลอน) — เว้นเป็นเส้นประให้ Maker กรอกเองด้วย Word
- template ที่ 2 (`subcommittee`) ที่สร้างในนี้เป็นแค่ตัวอย่างพิสูจน์ว่า multi-template ทำงานได้จริง
  (ต่างกันแค่คำว่า "คณะกรรมการบริษัท" → "คณะกรรมการชุดย่อย/อนุกรรมการ" ในหัวเรื่อง) — ผู้ใช้ปรับแก้/
  แต่ง layout เพิ่มเองได้ตรงๆด้วย Word ตามที่อธิบายไว้ด้านบน
"""
import os

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates")

FONT_NAME = "TH SarabunPSK"
FONT_SIZE = Pt(15)


def _set_run_font(run, bold: bool = False, italic: bool = False) -> None:
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.bold = bold
    run.italic = italic
    # ต้องตั้ง East Asian/complex-script font element ตรงๆด้วย ไม่งั้น Word อาจ fallback เป็นฟอนต์
    # อื่นสำหรับตัวอักษรไทย (python-docx's font.name ตั้งแค่ latin typeface element เฉยๆ)
    from docx.oxml.ns import qn

    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts")
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT_NAME)
    rfonts.set(qn("w:cs"), FONT_NAME)


def _p(doc, text: str, *, align=None, bold: bool = False, italic: bool = False):
    """เพิ่มพารากราฟข้อความล้วน 1 บรรทัด (run เดียว) — ใช้กับข้อความ static ที่ไม่มี Jinja tag ปน"""
    para = doc.add_paragraph()
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    _set_run_font(run, bold=bold, italic=italic)
    return para


def _tag_p(doc, tag_text: str):
    """พารากราฟที่มีแค่ docxtpl `{%p ... %}` tag ล้วนๆ (สำหรับ for/endfor/if ที่ต้องลบพารากราฟทิ้ง
    หลัง render) — ต้องเป็น run เดียวเป๊ะ ไม่ปนกับข้อความอื่น"""
    para = doc.add_paragraph()
    run = para.add_run(tag_text)
    _set_run_font(run)
    return para


def _build_document(doc_title_prefix: str) -> "docx.Document":
    """แกนกลางที่ template ทุกตัวใช้ร่วมกัน (layout/ฟอนต์/Jinja tag เหมือนกันหมด) — ต่างกันแค่
    `doc_title_prefix` (ข้อความก่อน "ครั้งที่ {{ meeting_number }}" เช่น "รายงานการประชุมคณะกรรมการ
    บริษัท" หรือ "รายงานการประชุมคณะกรรมการชุดย่อย") ให้ต่าง template สื่อสารประเภทการประชุมต่างกันได้
    โดยไม่ต้องเปลี่ยนโครงสร้าง/Jinja tag เลย"""
    doc = docx.Document()

    # ── Page setup: เลียนแบบไฟล์จริง (A4, margin ~1 นิ้ว, bottom แคบกว่าเล็กน้อย) ──────────
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.886)

    # ตั้ง default font ของ Normal style ด้วย (กันพารากราฟที่ python-docx สร้างเองเช่น table cell)
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = FONT_SIZE

    # ── หัวเอกสาร (จัดกึ่งกลาง เหมือนต้นฉบับ) ────────────────────────────────────────────
    _p(doc, "{{ company_name }}", align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, "{{ company_address }}", align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, "_______________________", align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, f"{doc_title_prefix} ครั้งที่ {{{{ meeting_number }}}}", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    doc.add_paragraph()

    # ── เวลาและสถานที่ ───────────────────────────────────────────────────────────────────
    _p(doc, "เวลาและสถานที่", bold=True)
    _p(
        doc,
        "ประชุมเมื่อวันที่ {{ meeting_date_thai }} ณ ห้องประชุม {{ company_name }} "
        "(“บริษัทฯ”) {{ company_address }}",
    )
    doc.add_paragraph()

    # ── ผู้เข้าร่วมประชุม (loop ผ่าน attendees — {%p %} ลบพารากราฟ tag ทิ้งหลัง render) ────
    _p(doc, "ผู้เข้าร่วมการประชุม", bold=True)
    _tag_p(doc, "{%p for attendee in attendees %}")
    para = doc.add_paragraph()
    r1 = para.add_run("{{ loop.index }}.\t")
    _set_run_font(r1)
    r2 = para.add_run("{{ attendee.name }}\t")
    _set_run_font(r2)
    r3 = para.add_run("{{ attendee.position }}")
    _set_run_font(r3)
    _tag_p(doc, "{%p endfor %}")
    doc.add_paragraph()

    # ── เปิดประชุม (ไม่ใส่เลของค์ประชุม/ร้อยละ — Module 3 ไม่สร้างตัวเลขนี้ กันหลอน) ────────
    _p(
        doc,
        "{{ chairperson_name }} ทำหน้าที่เป็นประธานในที่ประชุม (“ประธานฯ”) "
        "กล่าวเปิดการประชุม และดำเนินการประชุมตามระเบียบวาระดังต่อไปนี้",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    doc.add_paragraph()

    # ── วาระการประชุม (loop ผ่าน agenda_items) ──────────────────────────────────────────
    # เลขวาระ (2026-08-07, ผู้ใช้ขอ — ดู models.py's MeetingAgendaItem.label docstring): เดิม
    # hardcode prefix "วาระที่ {{ item.agenda_order }}" ตรงนี้ (index อัตโนมัติ 0,1,2,... เรียง
    # ต่อเนื่องเสมอ ไม่รองรับวาระย่อยแบบ 3.1/3.2 หรือเลขข้ามที่ไม่เรียงต่อเนื่องตามธรรมเนียมบอร์ดจริง)
    # เปลี่ยนเป็นพิมพ์ `{{ item.label }}` ตรงๆ ไม่เติม prefix อะไรให้เองอีกแล้ว — ผู้ใช้พิมพ์เต็มข้อความ
    # เองในช่อง label ได้อิสระ (ดีฟอลต์ที่ backend เติมให้ถ้าไม่กรอกคือ "วาระที่ {ลำดับ}" อยู่แล้ว ดู
    # main.py's _build_agenda_items() — หน้าตาเหมือนเดิมทุกประการถ้าไม่ได้ตั้งใจแก้เอง)
    _tag_p(doc, "{%p for item in agenda_items %}")
    para = doc.add_paragraph()
    r1 = para.add_run("{{ item.label }}\t")
    _set_run_font(r1, bold=True)
    r2 = para.add_run("{{ item.description }}")
    _set_run_font(r2, bold=True)

    _p(doc, "{{ item.discussion_summary }}", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    para = doc.add_paragraph()
    r1 = para.add_run("มติที่ประชุม\t")
    _set_run_font(r1, bold=True)
    r2 = para.add_run("({{ item.resolution_status_label }}) {{ item.resolution_text }}")
    _set_run_font(r2)

    _p(
        doc,
        "[หากมีรายละเอียดตาราง/ตัวเลขธุรกรรม/สัดส่วนหุ้นเพิ่มเติมสำหรับวาระนี้ "
        "กรุณาเพิ่มด้วยตนเองในไฟล์ Word นี้ก่อนส่งขออนุมัติ — AI ไม่สร้างตัวเลขธุรกรรมให้อัตโนมัติ]",
        italic=True,
    )
    doc.add_paragraph()
    _tag_p(doc, "{%p endfor %}")

    # ── เรื่องอื่นๆ ───────────────────────────────────────────────────────────────────────
    _p(doc, "เรื่องอื่นๆ", bold=True)
    _p(doc, "{{ other_business_notes }}", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    doc.add_paragraph()

    # ── ปิดประชุม + ลงชื่อ (เว้นเวลาปิดประชุม/ชื่อผู้บันทึกให้ Maker กรอกเอง — ไม่มีใน schema) ──
    _p(
        doc,
        "ประธานฯ กล่าวขอบคุณกรรมการทุกท่านที่เข้าร่วมประชุม และกล่าวปิดการประชุมในเวลา "
        "....................... น.",
    )
    doc.add_paragraph()
    doc.add_paragraph()
    _p(doc, "ลงชื่อ…………………………………………………………. ประธานในที่ประชุม")
    _p(doc, "                  ({{ chairperson_name }})")
    doc.add_paragraph()
    doc.add_paragraph()
    _p(doc, "ลงชื่อ…………………………………………………………. ผู้บันทึกการประชุม")
    _p(doc, "               (....................................................)")
    _p(doc, "                        เลขานุการบริษัท")
    doc.add_paragraph()

    # ── Footer เล็กๆ บอกที่มาเอกสาร (ตรวจสอบได้ว่า draft นี้ AI ช่วยร่าง ไม่ใช่ของมนุษย์ล้วน) ──
    _p(
        doc,
        "ร่างโดยระบบ AI ({{ generated_by_model }}) เมื่อ {{ generated_at_thai }} — "
        "ต้องผ่านการตรวจสอบและเพิ่มรายละเอียดตาราง/ตัวเลขธุรกรรมโดย Company Secretary "
        "ก่อนนำไปใช้จริง",
        italic=True,
    )
    return doc


def build_bod_minutes() -> str:
    """template หลัก (เดิม, ไม่เปลี่ยนชื่อไฟล์ — meeting ที่สร้างไว้ก่อนหน้านี้อ้างอิงชื่อนี้อยู่)"""
    doc = _build_document("รายงานการประชุมคณะกรรมการบริษัท")
    os.makedirs(TEMPLATE_DIR, exist_ok=True)
    path = os.path.join(TEMPLATE_DIR, "minutes_template.docx")
    doc.save(path)
    return path


def build_subcommittee() -> str:
    """template ตัวอย่างที่ 2 (ใหม่) — พิสูจน์ว่า multi-template ทำงานได้จริง ต่างจากตัวหลักแค่คำใน
    หัวเรื่อง ผู้ใช้ปรับแก้ layout/ถ้อยคำเพิ่มเองได้ตรงๆด้วย Word ทีหลัง (ดู docstring หัวไฟล์)"""
    doc = _build_document("รายงานการประชุมคณะกรรมการชุดย่อย/อนุกรรมการ")
    os.makedirs(TEMPLATE_DIR, exist_ok=True)
    path = os.path.join(TEMPLATE_DIR, "minutes_template_subcommittee.docx")
    doc.save(path)
    return path


if __name__ == "__main__":
    for build_fn in (build_bod_minutes, build_subcommittee):
        path = build_fn()
        print(f"เขียน template เสร็จแล้วที่: {path}")
